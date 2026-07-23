# -*- coding: utf-8 -*-
"""将 docs 下四份项目文档并入课程设计报告（满足提交要求第 3 项）。"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(r"c:\Users\ohmyc\Desktop\专业课程设计")
REPORT = ROOT / "docs" / "专业方向课程设计报告.docx"
OUT_DL = Path(r"c:\Users\ohmyc\Downloads") / "专业方向课程设计报告.docx"

DOCS = [
    ("附录 B  系统设计", ROOT / "docs" / "系统设计.md"),
    ("附录 C  功能说明", ROOT / "docs" / "功能说明.md"),
    ("附录 D  数据结构设计", ROOT / "docs" / "数据结构设计.md"),
    ("附录 E  API接口文档", ROOT / "docs" / "API接口文档.md"),
]


def set_run_font(run, *, size_pt=12, bold=False, east_asia="宋体"):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), east_asia)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.clear()
    run = p.add_run(text)
    size = {1: 16, 2: 14, 3: 12}.get(level, 12)
    set_run_font(run, size_pt=size, bold=True, east_asia="黑体")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5


def add_body(doc, text, *, first_line=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size_pt=size)


def strip_md(line: str) -> str:
    line = re.sub(r"^#+\s*", "", line)
    line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
    line = re.sub(r"`([^`]+)`", r"\1", line)
    line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
    return line.strip()


def append_markdown(doc: Document, title: str, md_path: Path):
    add_heading(doc, title, 1)
    add_body(
        doc,
        f"本附录与仓库文件 {md_path.name} 对应，满足课程设计提交对「项目文档」的要求。",
        first_line=True,
    )
    text = md_path.read_text(encoding="utf-8")
    in_code = False
    code_buf: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_body(doc, "\n".join(code_buf), first_line=False, size=10.5)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("|") and "---" in line:
            continue
        if line.startswith("#"):
            level = len(re.match(r"^#+", line).group(0))
            add_heading(doc, strip_md(line), min(level + 1, 3))
            continue
        if line.startswith("- ") or line.startswith("* "):
            add_body(doc, "• " + strip_md(line[2:]), first_line=False)
            continue
        if re.match(r"^\d+\.\s", line):
            add_body(doc, strip_md(line), first_line=False)
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            add_body(doc, " | ".join(cells), first_line=False, size=10.5)
            continue
        add_body(doc, strip_md(line), first_line=True)


def remove_old_appendices(doc: Document):
    """删除此前已追加的附录 B–E，避免重复生成。"""
    markers = ("附录 B", "附录 C", "附录 D", "附录 E")
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and "Heading" in p.style.name and any(p.text.startswith(m) for m in markers):
            start = i
            break
    if start is None:
        return
    for p in list(doc.paragraphs)[start:]:
        el = p._element
        el.getparent().remove(el)


def ensure_toc_entries(doc: Document):
    extras = [
        "附录 B  系统设计",
        "附录 C  功能说明",
        "附录 D  数据结构设计",
        "附录 E  API接口文档",
    ]
    existing = "\n".join(p.text for p in doc.paragraphs if p.style and str(p.style.name).startswith("toc"))
    if "附录 B" in existing:
        return
    # 插到最后一个 toc 之后
    last_toc = None
    for p in doc.paragraphs:
        if p.style and str(p.style.name).startswith("toc"):
            last_toc = p
    if last_toc is None:
        return
    parent = last_toc._element.getparent()
    idx = parent.index(last_toc._element) + 1
    created = []
    for text in extras:
        try:
            np = doc.add_paragraph(style="toc 1")
        except KeyError:
            np = doc.add_paragraph()
        np.clear()
        run = np.add_run(text)
        set_run_font(run, size_pt=12)
        created.append(np._element)
    for el in created:
        parent.remove(el)
    for i, el in enumerate(created):
        parent.insert(idx + i, el)


def main():
    if not REPORT.exists():
        raise SystemExit(f"报告不存在: {REPORT}")
    doc = Document(str(REPORT))
    remove_old_appendices(doc)
    for title, path in DOCS:
        if not path.exists():
            raise SystemExit(f"缺少文档: {path}")
        append_markdown(doc, title, path)
    ensure_toc_entries(doc)
    doc.save(str(REPORT))
    doc.save(str(OUT_DL))
    print(f"updated: {REPORT}")
    print(f"copied:  {OUT_DL}")


if __name__ == "__main__":
    main()
